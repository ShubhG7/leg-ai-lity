"""
Document parsing utilities for extracting placeholders from .docx files.
"""

import re
from docx import Document
from typing import List, Set
from .gemini_client import extract_placeholders_with_ai
from .placeholder_cleaner import clean_and_standardize_placeholders

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text content from a .docx file.
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
        
    except Exception as e:
        raise Exception(f"Error reading document: {str(e)}")

def extract_placeholders_regex(text: str) -> Set[str]:
    """
    Extract placeholders using precise regex patterns for legal documents.
    Focuses on proper legal document formatting.
    """
    placeholders = set()
    
    # Pattern 1: Standard legal brackets [Field Name]
    # Only capture if it looks like a proper legal field
    bracket_pattern = r'\[([A-Za-z\s]{2,50})\]'
    bracket_matches = re.findall(bracket_pattern, text)
    
    # Filter out informal text - only keep proper legal field names
    legal_field_keywords = [
        'name', 'amount', 'date', 'address', 'signature', 'company', 'investor',
        'purchase', 'valuation', 'cap', 'discount', 'rate', 'shares', 'stock',
        'equity', 'investment', 'funding', 'round', 'closing', 'effective',
        'governing', 'law', 'state', 'notice', 'email', 'phone', 'contact'
    ]
    
    for match in bracket_matches:
        match_lower = match.lower().strip()
        # Only include if it contains legal keywords or is clearly a field
        if (any(keyword in match_lower for keyword in legal_field_keywords) or
            len(match_lower.split()) <= 3):  # Short, likely field names
            placeholders.add(match.strip())
    
    # Pattern 2: Dollar amount placeholders like $[Amount]
    dollar_pattern = r'\$\[([A-Za-z\s]{2,30})\]'
    dollar_matches = re.findall(dollar_pattern, text)
    for match in dollar_matches:
        if any(keyword in match.lower() for keyword in ['amount', 'price', 'value', 'cost', 'investment', 'purchase']):
            placeholders.add(match.strip())
    
    # Pattern 3: Multiple underscores with legal context
    underscore_pattern = r'_{4,}'
    for match in re.finditer(underscore_pattern, text):
        # Get context around the underscores (50 chars before/after)
        start = max(0, match.start() - 50)
        end = min(len(text), match.end() + 50)
        context = text[start:end].lower()
        
        # Only include if context suggests it's a legal field
        if any(keyword in context for keyword in legal_field_keywords):
            # Try to extract a meaningful field name from context
            before_context = text[max(0, match.start() - 20):match.start()].strip()
            after_context = text[match.end():match.end() + 20].strip()
            
            # Look for field indicators
            field_indicators = ['name:', 'date:', 'amount:', 'address:', 'signature:']
            for indicator in field_indicators:
                if indicator in context:
                    field_name = indicator.replace(':', '').title()
                    placeholders.add(field_name)
                    break
            else:
                # Default to generic field name
                placeholders.add("Field")
    
    # Pattern 4: Explicit legal document placeholders
    # Look for patterns like "the [Field]" or "[Field] of the Company"
    legal_context_patterns = [
        r'(?:the\s+)?\[([A-Za-z\s]{2,30})\]',
        r'\[([A-Za-z\s]{2,30})\]\s+(?:of\s+the\s+)?(?:Company|Investor|Agreement)',
        r'(?:Company|Investor|Agreement)\s+\[([A-Za-z\s]{2,30})\]'
    ]
    
    for pattern in legal_context_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if any(keyword in match.lower() for keyword in legal_field_keywords):
                placeholders.add(match.strip())
    
    # Filter out any remaining informal or inappropriate text
    filtered_placeholders = set()
    for placeholder in placeholders:
        placeholder_lower = placeholder.lower().strip()
        
        # Skip if it contains informal language
        informal_words = ['oats', 'monday', 'next week', 'keep', 'you can', 'informal', 'title', 'field']
        if any(word in placeholder_lower for word in informal_words):
            continue
            
        # Skip generic terms (but allow them if they're part of a compound term)
        if len(placeholder.split()) == 1 and placeholder_lower in ['field', 'name', 'company', 'title', 'blank']:
            continue
            
        # Clean up the placeholder
        cleaned = placeholder.strip().title()
        if 3 <= len(cleaned) <= 50:  # Reasonable length, minimum 3 chars
            filtered_placeholders.add(cleaned)
    
    # Remove duplicates and merge similar ones
    final_placeholders = set()
    placeholder_list = list(filtered_placeholders)
    
    for placeholder in placeholder_list:
        # Check if this is a duplicate of an existing one
        is_duplicate = False
        for existing in final_placeholders:
            # Normalize for comparison (lowercase, no extra spaces)
            norm_placeholder = ' '.join(placeholder.lower().split())
            norm_existing = ' '.join(existing.lower().split())
            
            if (norm_placeholder == norm_existing or 
                norm_placeholder in norm_existing or 
                norm_existing in norm_placeholder):
                # Keep the one with better formatting (proper capitalization)
                if placeholder.count(' ') > existing.count(' '):  # More descriptive
                    final_placeholders.discard(existing)
                    final_placeholders.add(placeholder)
                elif placeholder.count(' ') == existing.count(' ') and len(placeholder) > len(existing):
                    final_placeholders.discard(existing)
                    final_placeholders.add(placeholder)
                is_duplicate = True
                break
        
        if not is_duplicate:
            final_placeholders.add(placeholder)
    
    # Use the advanced cleaner to standardize placeholders
    return clean_and_standardize_placeholders(final_placeholders)

async def extract_all_placeholders(file_path: str) -> List[str]:
    """
    Extract placeholders using Gemini AI ONLY for maximum accuracy.
    """
    # Extract text from document
    document_text = extract_text_from_docx(file_path)
    
    print(f"🤖 Using Gemini AI for placeholder detection on {len(document_text)} character document...")
    
    # Use Gemini AI - trust its results completely
    ai_placeholders = await extract_placeholders_with_ai(document_text)
    
    if len(ai_placeholders) > 0:
        print(f"✅ Gemini AI found {len(ai_placeholders)} placeholders")
        # Remove duplicates but don't filter out any results
        unique_placeholders = list(dict.fromkeys(ai_placeholders))  # Preserve order, remove duplicates
        print(f"📋 After deduplication: {len(unique_placeholders)} placeholders")
        print(f"📝 Placeholders: {unique_placeholders}")
        return unique_placeholders
    else:
        print("❌ Gemini AI failed, falling back to regex...")
        # Fallback to regex if AI fails
        regex_placeholders = extract_placeholders_regex(document_text)
        return list(set(regex_placeholders))
