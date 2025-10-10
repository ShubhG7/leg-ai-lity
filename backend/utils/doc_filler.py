"""
Document filling utilities for replacing placeholders with user data.
"""

import re
from docx import Document
from typing import Dict
from typing import Optional
from datetime import datetime
import re as _re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import re
import os
import tempfile


"""
Document filling utilities for replacing placeholders with user data.
"""


def _normalize_money(raw: str) -> str:
    s = (raw or "").strip()
    if not s:
        return s
    s_lower = s.lower().replace(" ", "")
    # 8m, 8M, 8 million, 8,000,000, $8000000
    m = _re.match(r"^\$?([0-9][0-9,]*)(\.[0-9]+)?$", s.replace(" ", ""))
    if m:
        num = float((m.group(1) + (m.group(2) or "")).replace(",", ""))
        return f"${num:,.2f}".rstrip("0").rstrip(".")
    m = _re.match(r"^\$?([0-9]+)(m|million)$", s_lower)
    if m:
        num = float(m.group(1)) * 1_000_000
        return f"${num:,.0f}"
    m = _re.match(r"^\$?([0-9]+)(k|thousand)$", s_lower)
    if m:
        num = float(m.group(1)) * 1_000
        return f"${num:,.0f}"
    # Already starts with $
    if s.startswith("$"):
        return s
    return s


def _normalize_date(raw: str) -> str:
    s = (raw or "").strip()
    if not s:
        return s
    # Try multiple common formats without extra deps
    candidates = [
        "%B %d, %Y", "%b %d, %Y", "%d %B %Y", "%d %b %Y",
        "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y",
    ]
    for fmt in candidates:
        try:
            dt = datetime.strptime(s, fmt)
            return dt.strftime("%B %d, %Y")
        except Exception:
            pass
    # Simple natural language like "Oct 9 2025" or "October 9 2025"
    try:
        # Insert comma if "Month D YYYY"
        m = _re.match(r"^(\w+)\s+(\d{1,2}),?\s+(\d{4})$", s)
        if m:
            dt = datetime.strptime(f"{m.group(1)} {m.group(2)} {m.group(3)}", "%B %d %Y")
            return dt.strftime("%B %d, %Y")
    except Exception:
        pass
    return s


def _parse_date(raw: str) -> Optional[datetime]:
    s = (raw or "").strip()
    if not s:
        return None
    fmts = [
        "%B %d, %Y", "%b %d, %Y", "%d %B %Y", "%d %b %Y",
        "%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%m/%d/%y", "%d/%m/%y",
    ]
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt)
        except Exception:
            continue
    # Try Month D YYYY without comma
    try:
        m = _re.match(r"^(\w+)\s+(\d{1,2}),?\s+(\d{2,4})$", s)
        if m:
            year = m.group(3)
            if len(year) == 2:
                year = "20" + year
            return datetime.strptime(f"{m.group(1)} {m.group(2)} {year}", "%B %d %Y")
    except Exception:
        pass
    return None


def _normalize_email(raw: str) -> str:
    s = (raw or "").strip()
    if _re.match(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$", s):
        return s.lower()
    return s


def _title_case_name(raw: str) -> str:
    s = (raw or "").strip()
    return " ".join([w.capitalize() if w else w for w in s.split()])


def normalize_placeholder_values(data: Dict[str, str]) -> Dict[str, str]:
    """
    Normalize user-provided values based on placeholder names.
    """
    normalized: Dict[str, str] = {}
    for key, value in (data or {}).items():
        v = (value or "").strip().strip('\"\'')
        kl = key.lower()
        if any(x in kl for x in ["amount", "price", "valuation", "purchase", "cap", "investment"]):
            v = _normalize_money(v)
        elif any(x in kl for x in ["date", "day", "closing", "effective"]):
            v = _normalize_date(v)
        elif "email" in kl:
            v = _normalize_email(v)
        elif any(x in kl for x in ["name", "investor", "company", "title"]):
            v = _title_case_name(v)
        normalized[key] = v
    return normalized


def fill_document_placeholders(input_file_path: str, output_file_path: str, placeholder_data: Dict[str, str]) -> str:
    """
    Fill placeholders in a .docx document with provided data using robust, formatting-preserving methods.
    Strategy:
      1) If MERGEFIELDs exist, use MailMerge to merge values.
      2) Else, convert bracket placeholders to Jinja and render via docxtpl.
      3) Afterwards, fill content controls (SDTs) and underscores contextually.
    """
    try:
        data_normalized = normalize_placeholder_values(placeholder_data)
        if _doc_has_merge_fields(input_file_path):
            _fill_with_mailmerge(input_file_path, output_file_path, data_normalized)
            _fill_content_controls_inplace(output_file_path, data_normalized)
            _fill_underscores_inplace(output_file_path, data_normalized)
            return output_file_path
        with tempfile.TemporaryDirectory() as tmpdir:
            transformed_tpl_path = os.path.join(tmpdir, "template_transformed.docx")
            _transform_doc_to_jinja(input_file_path, transformed_tpl_path, data_normalized)
            _render_jinja_template(transformed_tpl_path, output_file_path, data_normalized)
        _fill_content_controls_inplace(output_file_path, data_normalized)
        _fill_underscores_inplace(output_file_path, data_normalized)
        return output_file_path
    except Exception as e:
        raise Exception(f"Error filling document: {str(e)}")


def _doc_has_merge_fields(docx_path: str) -> bool:
    try:
        from mailmerge import MailMerge  # type: ignore
        with MailMerge(docx_path) as mm:  # type: ignore
            fields = mm.get_merge_fields()
            return bool(fields)
    except Exception:
        return False


essential_money_keys = ["amount", "price", "valuation", "purchase", "cap", "investment"]
essential_date_keys = ["date", "day", "closing", "effective"]


def _fill_with_mailmerge(input_file_path: str, output_file_path: str, data_normalized: Dict[str, str]) -> None:
    from mailmerge import MailMerge  # type: ignore

    def normalize_key(s: str) -> str:
        return _re.sub(r"[^a-z0-9]", "", (s or "").lower())

    with MailMerge(input_file_path) as mm:  # type: ignore
        merge_fields = list(mm.get_merge_fields())
        mapping: Dict[str, str] = {}
        for field in merge_fields:
            key_norm = normalize_key(field)
            val: Optional[str] = None
            for k, v in (data_normalized or {}).items():
                if normalize_key(k) == key_norm:
                    val = v
                    break
            field_lower = field.lower()
            if val is not None:
                if any(x in field_lower for x in essential_money_keys):
                    val = _normalize_money(val)
                elif any(x in field_lower for x in essential_date_keys):
                    parsed = _parse_date(val)
                    val = parsed.strftime("%B %d, %Y") if parsed else _normalize_date(val)
            if val is not None:
                mapping[field] = val
        if mapping:
            mm.merge(**mapping)
        mm.write(output_file_path)


def _transform_doc_to_jinja(input_path: str, output_path: str, data_normalized: Dict[str, str]) -> None:
    doc = Document(input_path)

    def convert_text(text: str) -> str:
        if not text:
            return text

        def norm_label(s: str) -> str:
            return _re.sub(r"\s+", " ", (s or "").strip())

        def to_jinja_var(label: str) -> str:
            # Convert arbitrary label to a safe Jinja identifier: snake_case
            base = _re.sub(r"\s+", " ", (label or "").strip())
            snake = _re.sub(r"[^A-Za-z0-9]+", "_", base).strip("_")
            if not snake:
                snake = "field"
            if snake[0].isdigit():
                snake = f"f_{snake}"
            return snake.lower()

        def bracket_repl(m: re.Match[str]) -> str:
            inside = m.group(1).strip()
            jvar = to_jinja_var(inside)
            return f"{{{{ {jvar} }}}}"

        def dollar_repl(m: re.Match[str]) -> str:
            inside = m.group(1).strip()
            jvar = to_jinja_var(inside)
            return f"{{{{ {jvar} }}}}"

        text2 = re.sub(r"\$\[([^\]]+)\]", dollar_repl, text)
        text2 = re.sub(r"\[([^\]]+)\]", bracket_repl, text2)
        return text2

    def convert_paragraph(paragraph: Paragraph) -> None:
        original = paragraph.text
        converted = convert_text(original)
        if converted != original:
            paragraph.clear()
            paragraph.add_run(converted)

    def iter_all_paragraphs(d: Document):
        for p in d.paragraphs:
            yield p
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        yield p
        for s in d.sections:
            for p in s.header.paragraphs:
                yield p
            for t in s.header.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p
            for p in s.footer.paragraphs:
                yield p
            for t in s.footer.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            yield p

    for p in iter_all_paragraphs(doc):
        convert_paragraph(p)

    doc.save(output_path)


def _render_jinja_template(template_path: str, output_path: str, context: Dict[str, str]) -> None:
    from docxtpl import DocxTemplate  # type: ignore

    tpl = DocxTemplate(template_path)

    # Filters optional; values are pre-normalized
    jenv = getattr(tpl, "jinja_env", None)
    if jenv is not None:
        def money_filter(val: Optional[str]) -> str:
            return _normalize_money(val or "")
        def date_filter(val: Optional[str], fmt: str = "%B %d, %Y") -> str:
            if not val:
                return ""
            dt = _parse_date(val)
            return dt.strftime(fmt) if dt else _normalize_date(val)
        def titlecase_filter(val: Optional[str]) -> str:
            return _title_case_name(val or "")
        def loweremail_filter(val: Optional[str]) -> str:
            return _normalize_email(val or "")
        jenv.filters["money"] = money_filter
        jenv.filters["datefmt"] = date_filter
        jenv.filters["titlecase"] = titlecase_filter
        jenv.filters["loweremail"] = loweremail_filter

    def to_jinja_var(label: str) -> str:
        base = _re.sub(r"\s+", " ", (label or "").strip())
        snake = _re.sub(r"[^A-Za-z0-9]+", "_", base).strip("_")
        if not snake:
            snake = "field"
        if snake[0].isdigit():
            snake = f"f_{snake}"
        return snake.lower()

    render_ctx: Dict[str, str] = {}
    for k, v in (context or {}).items():
        render_ctx[to_jinja_var(k)] = v

    tpl.render(render_ctx)
    tpl.save(output_path)


def _fill_content_controls_inplace(docx_path: str, data_normalized: Dict[str, str]) -> None:
    doc = Document(docx_path)
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    VAL_ATTR = "{" + W_NS + "}val"

    def normalize_key(s: str) -> str:
        return _re.sub(r"[^a-z0-9]", "", (s or "").lower())

    key_map = {normalize_key(k): v for k, v in (data_normalized or {}).items()}

    def set_sdt_text(sdt) -> None:
        # Find alias/tag without passing namespaces
        alias_nodes = sdt.xpath('.//*[local-name()="sdtPr"]/*[local-name()="alias"]')
        tag_nodes = sdt.xpath('.//*[local-name()="sdtPr"]/*[local-name()="tag"]')
        title = None
        if alias_nodes:
            title = alias_nodes[0].get(VAL_ATTR) or title
        if not title and tag_nodes:
            title = tag_nodes[0].get(VAL_ATTR)
        if not title:
            return
        key = normalize_key(title)
        val = key_map.get(key)
        if val is None:
            return
        content_nodes = sdt.xpath('.//*[local-name()="sdtContent"]')
        if not content_nodes:
            return
        sdt_content = content_nodes[0]
        for child in list(sdt_content):
            sdt_content.remove(child)
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = val
        r.append(t)
        p.append(r)
        sdt_content.append(p)

    def process_element(elem) -> None:
        sdts = elem.xpath('.//*[local-name()="sdt"]')
        for sdt in sdts:
            set_sdt_text(sdt)

    process_element(doc.element.body)
    for s in doc.sections:
        process_element(s.header._element)
        process_element(s.footer._element)
    doc.save(docx_path)


def _replace_underscore_placeholders(doc: Document, placeholder_data: Dict[str, str]):
    """
    Replace underscore placeholders (___) with appropriate values based on context.
    - Specifically handles signature blocks like "By:", "Name:", "Title:", "Date:", "Address:", "Email:" with party bias (Company/Investor).
    - Also applies general contextual replacement as a fallback.
    """
    underscore_pattern = r'_{3,}'
    labeled_pattern = r'(By|Name|Title|Date|Address|Email)\s*:\s*_{3,}'
    bare_label_line_pattern = r'^(.*?\b(By|Name|Title|Date|Address|Email)\s*:)\s*$'

    def normalize_key(s: str) -> str:
        return _re.sub(r"[^a-z0-9]", "", (s or "").lower())

    def get_party_context(text: str) -> Optional[str]:
        tl = (text or "").lower()
        if 'company' in tl and 'investor' in tl:
            ci = tl.find('company')
            ii = tl.find('investor')
            return 'company' if ci <= ii else 'investor'
        if 'company' in tl:
            return 'company'
        if 'investor' in tl:
            return 'investor'
        return None

    def choose_value_for_label(label: str, party: Optional[str]) -> Optional[str]:
        label_l = (label or '').lower()
        candidates = list(placeholder_data.items())
        # Party-specific first
        if party:
            if label_l in ('by', 'name', 'signature'):
                for k, v in candidates:
                    kl = k.lower()
                    if party in kl and (('signature' in kl) or ('signatory' in kl) or ('name' in kl)):
                        return v
            if label_l == 'title':
                for k, v in candidates:
                    kl = k.lower()
                    if party in kl and ('title' in kl or 'designation' in kl or 'position' in kl):
                        return v
            if label_l == 'date':
                for k, v in candidates:
                    kl = k.lower()
                    if party in kl and (('signature date' in kl) or ('date' in kl)):
                        return v
            if label_l == 'address':
                for k, v in candidates:
                    kl = k.lower()
                    if party in kl and 'address' in kl:
                        return v
            if label_l == 'email':
                for k, v in candidates:
                    kl = k.lower()
                    if party in kl and 'email' in kl:
                        return v
        # Generic fallbacks
        if label_l in ('by', 'name', 'signature'):
            for k, v in candidates:
                kl = k.lower()
                if ('signature' in kl) or ('signatory' in kl) or ('name' in kl):
                    return v
        if label_l == 'title':
            for k, v in candidates:
                if 'title' in k.lower() or 'designation' in k.lower() or 'position' in k.lower():
                    return v
        if label_l == 'date':
            for k, v in candidates:
                if 'signature date' in k.lower():
                    return v
            for k, v in candidates:
                if 'date' in k.lower():
                    return v
        if label_l == 'address':
            for k, v in candidates:
                if 'address' in k.lower():
                    return v
        if label_l == 'email':
            for k, v in candidates:
                if 'email' in k.lower():
                    return v
        return None

    def replace_in_text(text: str) -> str:
        if not text:
            return text
        party = get_party_context(text)
        # Pass 1: labeled fields with underscores
        def repl_labeled(m: re.Match[str]) -> str:
            label = m.group(1)
            val = choose_value_for_label(label, party)
            replacement = val or ''
            return f"{label}: {replacement}"
        new_text = re.sub(labeled_pattern, repl_labeled, text)
        # Pass 2: bare label lines with no underscores (just ":")
        def repl_bare_label(m: re.Match[str]) -> str:
            prefix = m.group(1)
            label = m.group(2)
            val = choose_value_for_label(label, party)
            replacement = val or ''
            return f"{prefix} {replacement}"
        new_text = re.sub(bare_label_line_pattern, repl_bare_label, new_text)
        # Pass 3: remaining underscores -> general heuristic
        def resolve_value_for_context(context_text: str) -> Optional[str]:
            ct = (context_text or "").lower()
            best_key = None
            local_party = get_party_context(context_text) or party
            for k in placeholder_data.keys():
                kl = k.lower()
                if any(word in ct for word in kl.split()):
                    best_key = k
                    if local_party and local_party in kl:
                        return placeholder_data[k]
                    if any(x in kl for x in ["address", "email", "name", "amount", "date", "title", "signature"]):
                        return placeholder_data[k]
            return placeholder_data.get(best_key) if best_key else None
        while re.search(underscore_pattern, new_text):
            ctx_val = resolve_value_for_context(new_text)
            replacement = ctx_val or ""
            new_text = re.sub(underscore_pattern, replacement, new_text, count=1)
        return new_text

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if re.search(underscore_pattern, paragraph.text) or re.search(labeled_pattern, paragraph.text) or re.search(bare_label_line_pattern, paragraph.text):
                new_text = replace_in_text(paragraph.text)
                if new_text != paragraph.text:
                    paragraph.clear()
                    paragraph.add_run(new_text)

    # Body paragraphs and tables
    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # Headers and footers
    for s in doc.sections:
        process_paragraphs(s.header.paragraphs)
        for t in s.header.tables:
            for r in t.rows:
                for c in r.cells:
                    process_paragraphs(c.paragraphs)
        process_paragraphs(s.footer.paragraphs)
        for t in s.footer.tables:
            for r in t.rows:
                for c in r.cells:
                    process_paragraphs(c.paragraphs)


def _fill_underscores_inplace(docx_path: str, data: Dict[str, str]) -> None:
    doc = Document(docx_path)
    _replace_underscore_placeholders(doc, data)
    doc.save(docx_path)
